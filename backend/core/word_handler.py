"""Handler for DOCX (Word) files."""

import logging
from pathlib import Path
from typing import Tuple

logger = logging.getLogger(__name__)


class WordHandler:
    """Extract text from DOCX (Word) files."""

    SUPPORTED_EXTENSIONS = ('.docx',)

    @staticmethod
    def read_word(file_path: str) -> Tuple[str, str | None]:
        """
        Extract text from a DOCX file.

        Args:
            file_path: Path to the DOCX file

        Returns:
            (content, error) - content is str, error is None on success
        """
        try:
            from docx import Document
        except ImportError:
            logger.warning("python-docx not installed. Install with: pip install python-docx")
            return "", "python-docx library not installed"

        try:
            doc = Document(file_path)
            text_parts = []

            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    text_parts.append(" | ".join(row_text))

            content = "\n".join(text_parts)
            logger.info(f"Extracted {len(content)} characters from Word file: {Path(file_path).name}")
            return content, None

        except FileNotFoundError:
            error_msg = f"File not found: {file_path}"
            logger.error(error_msg)
            return "", error_msg
        except Exception as e:
            error_msg = f"Error reading Word file: {str(e)}"
            logger.error(error_msg)
            return "", error_msg
